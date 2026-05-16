from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Create a new Document
doc = Document()

# Title
title = doc.add_heading('HW #2 – Functional Unit and System Boundary', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Question 1
doc.add_heading('Question 1: Definition of Functional Unit and Its Criticality in LCA Studies', level=1)

p1 = doc.add_paragraph(
    'A functional unit is a standardized, measurable quantity that defines the service or function delivered by a product system. Rather than comparing products based solely on their physical properties, the functional unit establishes a common reference point that allows fair comparison between different products that serve the same purpose. In life cycle assessment, the functional unit is critical because it ensures that we are comparing systems on the basis of equivalent utility.'
)

p2 = doc.add_paragraph(
    'The functional unit is essential in LCA studies for several reasons. First, it enables meaningful comparative assertions between competing products or systems. Without a clear functional unit, comparisons become arbitrary—comparing one liter of milk to one kilogram of milk, for example, would yield different results even though they may represent vastly different amounts of the same product. Second, the functional unit anchors all inventory data and impact calculations, ensuring that environmental burdens are properly attributed to the service provided. Finally, the functional unit allows LCA results to be transparent and reproducible; future studies can use the same functional unit to validate or expand upon previous findings.'
)

# Question 2
doc.add_heading('Question 2: Definition of System Boundary', level=1)

p3 = doc.add_paragraph(
    'A system boundary is the imaginary line that separates the processes we choose to include in our life cycle assessment from those we deliberately exclude. The system boundary defines the scope of the analysis by specifying which stages of a product\'s lifecycle, and which activities or inputs within those stages, will be analyzed for their environmental impacts.'
)

p4 = doc.add_paragraph(
    'Establishing a clear system boundary is fundamental because it determines the completeness and relevance of the LCA. A narrow boundary might miss critical environmental hotspots; a boundary that is too broad becomes technically infeasible and loses focus. The system boundary is defined by both the lifecycle stages considered (from raw material extraction to end-of-life) and by the depth of analysis within each stage. For example, an LCA might include capital equipment energy in some boundaries but exclude it in others. Clear documentation of the system boundary is essential for interpreting LCA results correctly and for allowing others to assess the representativeness and appropriateness of the study.'
)

# Question 3
doc.add_heading('Question 3: Analysis of the Chosen System – Milk Production', level=1)

# 3.1
doc.add_heading('3.1 Functional Unit for the Milk System', level=2)

p5 = doc.add_paragraph(
    'The functional unit for this analysis is: '
)
p5_run = p5.add_run('One carton containing one liter of refrigerated milk, delivered to the consumer ready for consumption.')
p5_run.bold = True

p6 = doc.add_paragraph(
    'This functional unit captures the essential service being provided—delivering a specific volume of drinkable milk in its typical commercial packaging—while specifying that the milk is refrigerated, which is a material characteristic distinguishing the two systems being compared (both dairy and soy milk in their commercially available, refrigerated forms). This functional unit aligns with the standardized approach used in comprehensive life cycle assessment databases that enable fair comparison between dairy and plant-based milk alternatives (Poore & Nemecek, 2018).'
)

# 3.2
doc.add_heading('3.2 System Boundary Diagram and Lifecycle Stages', level=2)

doc.add_paragraph('[INSERT HAND-DRAWN DIAGRAM HERE: Cradle-to-grave system boundary showing three main stages with inputs and outputs]')
doc.add_paragraph('')

p7 = doc.add_paragraph(
    'The system boundary for milk production (both dairy and soy) encompasses the following six lifecycle stages:'
)

# Stage 1
doc.add_heading('Stage 1: Raw Material Extraction', level=3)

p8 = doc.add_paragraph(
    'For dairy milk, raw materials include cattle feed (forages such as alfalfa and ryegrass, grains such as corn and barley, and byproducts such as soybean meal and cotton seeds), and water. The water footprint of dairy milk production is substantial: approximately 628 liters of water are required per liter of milk produced, accounting for drinking water and water consumed in feed production (Poore & Nemecek, 2018). Additionally, this stage includes the energy required for farm operations.'
)

p9 = doc.add_paragraph(
    'For soy milk, raw materials consist of soybeans that are grown, harvested, and transported, water (approximately 27–28 liters per liter of soy milk produced), and minerals/additives (Poore & Nemecek, 2018). This represents approximately 22 times less water than dairy milk production. This stage also includes energy for crop cultivation.'
)

p10 = doc.add_paragraph(
    'A critical process often underestimated in lifecycle analyses is the '
)
p10_run = p10.add_run('transportation of inputs')
p10_run.bold = True
p10.add_run('. Feed for dairy cattle and soybeans for processing must be transported from farms to processing facilities, generating emissions from transportation fuel. This represents a significant but sometimes overlooked contribution to the overall system.')

# Stage 2
doc.add_heading('Stage 2: Material Processing & Manufacturing', level=3)

p11 = doc.add_paragraph(
    'For dairy milk, this stage involves pasteurization—heating milk to kill pathogens—followed by cooling to maintain freshness. Dairy processing is energy-intensive due to the requirement for continuous refrigeration throughout the process. Other processing steps include separation and filtration.'
)

p12 = doc.add_paragraph(
    'For soy milk, processing begins with soaking dried soybeans in water. The soaked beans are then ground mechanically into a slurry, which is brought to a boil (15–20 minutes) to inactivate soybean trypsin inhibitor and sterilize the product (Cao et al., 2009). Complete inactivation of trypsin inhibitor activity can be achieved through thermal processing at these temperatures and durations. Following boiling, insoluble residues (soy pulp) are removed by straining and filtration. Chemical additives, specifically defoaming agents (typically glycerin-based fatty acid esters), are used during the boiling step to manage foam formation (Geburt et al., 2022).'
)

# Stage 3
doc.add_heading('Stage 3: Assembly, Packaging, & Preparation', level=3)

p13 = doc.add_paragraph(
    'For both systems, the milk is packaged in a carton composed of paper, plastic (polyethylene), and aluminum foil layers. Carton materials must be extracted and processed before being used for packaging. The packaging process involves filling the carton, sealing it, and adding labels. Each of these sub-processes requires energy and contributes to the overall environmental footprint.'
)

# Stage 4
doc.add_heading('Stage 4: Distribution & Retail Delivery', level=3)

p14 = doc.add_paragraph(
    'Both refrigerated dairy milk and refrigerated soy milk require maintained cold chain logistics from the processing facility to the retail point. This involves refrigerated transport trucks, distribution center storage at controlled temperatures (33–38°F), and retail store refrigeration and display. The energy demand for maintaining these cold temperatures is substantial and represents a significant environmental impact for both systems.'
)

# Stage 5
doc.add_heading('Stage 5: Use Phase / Operation & Maintenance', level=3)

p15 = doc.add_paragraph(
    'At the consumer level, both milk types are stored in home refrigerators at similar temperatures. The milk is consumed, and the carton, along with any residual milk, is removed from the refrigerator.'
)

p16 = doc.add_paragraph(
    'The use phase differs slightly between products in terms of spoilage risk and shelf life. Properly refrigerated dairy milk maintains freshness for approximately 2–3 weeks. Refrigerated soy milk, when refrigerated, also maintains freshness for approximately 2–3 weeks. Both require continuous refrigeration; without it, they spoil rapidly due to microbial growth.'
)

# Stage 6
doc.add_heading('Stage 6: End-of-Life', level=3)

p17 = doc.add_paragraph(
    'Both milk cartons are recyclable. The end-of-life stage includes collecting cartons, sorting materials (separating paper, plastic, and aluminum), and processing them for reuse or energy recovery. Wastewater treatment of residual milk also falls within this stage.'
)

# 3.4
doc.add_heading('3.4 System Boundary Completeness Assessment', level=2)

p18 = doc.add_paragraph(
    'This analysis has attempted to include every major process within the defined system boundary. The primary environmental hotspot in this system is '
)
p18_run = p18.add_run('raw milk production')
p18_run.bold = True
p18.add_run('. According to a comprehensive systematic review of 84 dairy LCA studies published between 2018–2024, raw milk production consistently emerges as the largest environmental contributor across all impact categories, with fertilizer use, agricultural material production, and on-site emissions as the primary drivers for global warming potential, acidification, and eutrophication (González-García et al., 2025). For plant-based milk production, by contrast, the environmental hotspots shift to processing, packaging, and transportation stages rather than raw material production (Geburt et al., 2022).')

p18b = doc.add_paragraph('')

p19 = doc.add_paragraph(
    'One significant process was included that could reasonably be excluded: '
)
p19_run = p19.add_run('retail store refrigeration energy')
p19_run.bold = True
p19.add_run('. While retail refrigeration is part of the distribution system, some LCA studies exclude retailer-specific energy on the grounds that retailers\' refrigeration systems serve multiple products and attributing a specific share to any single product is complex and imprecise. In this study, retail refrigeration is included because both dairy and soy milk require identical cold-chain management at retail, and this shared requirement is material to understanding the comparison. The environmental impact of retail refrigeration is likely negligible relative to the major production impacts for each system, but it is justifiable to include for completeness.')

# Question 4
doc.add_heading('Question 4: Functional Units for Non-Selected Cases', level=1)

doc.add_heading('Your journey to the university (Bus vs. Private EV Car)', level=3)
p20 = doc.add_paragraph('Functional unit: One round-trip commute from Tel Aviv to Reichman University in Herzliya and back to Tel Aviv, representing a typical student journey to campus.')

doc.add_heading('Take-away coffee', level=3)
p21 = doc.add_paragraph('Functional unit: One 200-milliliter serving of coffee, delivered in a consumable form.')

doc.add_heading('Shopping bags', level=3)
p22 = doc.add_paragraph('Functional unit: One plastic shopping bag.')

doc.add_heading('Producing electricity', level=3)
p23 = doc.add_paragraph('Functional unit: One megawatt-hour (MWh) of electricity delivered to the grid.')

# References
doc.add_heading('References', level=1)

ref1 = doc.add_paragraph('Cao, X., Gu, Z., Zhang, Q., Zhang, S., & Wang, Z. (2009). Elimination of trypsin inhibitor activity and beany flavor in soy milk by consecutive blanching and ultrahigh-temperature (UHT) processing. ')
ref1_run = ref1.add_run('Journal of Agricultural and Food Chemistry')
ref1_run.italic = True
ref1.add_run(', 57(14), 6362-6368. https://doi.org/10.1021/jf801039h')

ref2 = doc.add_paragraph('Geburt, K., Albrecht, E. H., Pointke, M., Pawelzik, E., Gerken, M., & Traulsen, I. (2022). A comparative analysis of plant-based milk alternatives Part 2: Environmental impacts. ')
ref2_run = ref2.add_run('Sustainability')
ref2_run.italic = True
ref2.add_run(', 14(14), 8424. https://doi.org/10.3390/su14148424')

ref3 = doc.add_paragraph('González-García, S., Feijoo, G., & Moreira, M. T. (2025). Life cycle assessment applied to milk production and processing: An integrative systematic literature review. ')
ref3_run = ref3.add_run('Sustainability')
ref3_run.italic = True
ref3.add_run(', 17(4), 1615. https://doi.org/10.3390/su17041615')

ref4 = doc.add_paragraph('Poore, J., & Nemecek, T. (2018). Reducing food\'s environmental impacts through producers and consumers. ')
ref4_run = ref4.add_run('Science')
ref4_run.italic = True
ref4.add_run(', 360(6392), 987-992. https://doi.org/10.1126/science.aaq0216')

# Save document with updated filename to avoid lock issues
doc.save(r'c:\עמרי\Shmags 2\projects\university\notes\HW2_Functional_Unit_and_System_Boundary_UPDATED.docx')
print("Word document created successfully with updated citations!")
print("Saved as: HW2_Functional_Unit_and_System_Boundary_UPDATED.docx")
