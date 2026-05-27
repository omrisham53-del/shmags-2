from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading("T'horn Manor Library - Adventure 2 Session 1", 0)

# Session Goals
doc.add_heading("Session Goals", 1)
goals = [
    "Learn culture and historical knowledge about dragon riders",
    "Expose the role of blue salt crystals in bonding with dragons",
    "Herald discovers a new path aligned with dragon rider values",
    "Aerendil begins to realize their patron may be an ancient dragon",
    "Expose the Bard's past story and connection to T'horn",
    "Discover a map showing traces of a researcher in the library",
    "Establish dragon riders as a central campaign theme",
    "Introduce the exiled scholar as a potential guide outside the city"
]
for goal in goals:
    doc.add_paragraph(goal, style='List Bullet')

# Scene 1
doc.add_heading("Scene 1: T'horn Manor Library", 1)
doc.add_heading("Setting", 2)
doc.add_paragraph(
    "T'horn's library is a vast space with towering shelves reaching the ceiling. Natural light flows through "
    "large windows draped in heavy curtains. The floor is covered with stacks of books, maps, and remnants of "
    "enchanted items covered in dust. The scent of old paper and plaster fills the air. The library is an almost "
    "chaotic catalog, but for those who know how to search, it contains a treasure of knowledge."
)

doc.add_heading("Opportunities", 2)
opps = [
    "Three historical stories about dragon riders (old book, paintings, handwriting)",
    "Small blue crystals in a glass box (unclear why they're here)",
    "Low-level magical items (DC 16-20 or DC 20+ if they really try)",
    "Mundane items - books, maps, tins, coins, fragments",
    "Chance for the Bard to talk about his past if he wants",
    "T'horn's comment that he's somewhat escaping Guard duty (opening to the subject)"
]
for opp in opps:
    doc.add_paragraph(opp, style='List Bullet')

doc.add_heading("Character-Specific Elements", 2)

p = doc.add_paragraph()
p.add_run("Herald: ").bold = True
p.add_run(
    "Should be exposed to the heroic story of dragon riders. Should see the difference between dragon rider values "
    "(freedom, partnership) and the Order that replaced them. This should begin to stir an inner voice: 'This is "
    "something I could be.' Give him quiet moments (alone in parts of the library) to raise these feelings."
)

p = doc.add_paragraph()
p.add_run("Aerendil: ").bold = True
p.add_run(
    "Should discover the story of the fiery legacy of dragon riders. Should begin to realize that his mysterious "
    "patron may be an ancient dragon. Should begin to translate the red/fire dreams as part of this pact. He doesn't "
    "need to understand everything yet - there are still mysteries."
)

p = doc.add_paragraph()
p.add_run("Bard: ").bold = True
p.add_run(
    "Will be relatively quiet during the library. At some point (after T'horn reveals his deal), he can share about "
    "his past and previous band. This should be a group moment."
)

doc.add_heading("DM Notes", 2)
notes = [
    "If they want to explore something, give them freedom to discover",
    "If they seem lost, give small hints",
    "Save the map for the end of the session",
    "When they find the blue crystals - that's a small hint they're connected to dragon riders"
]
for note in notes:
    doc.add_paragraph(note, style='List Bullet')

# Scene 2
doc.add_heading("Scene 2: Meeting T'horn", 1)
doc.add_heading("Setting", 2)
doc.add_paragraph(
    "T'horn is a man in his 50s-60s, full of strength but somewhat disheveled in appearance. He likes to sound "
    "thoughtful but also melancholic - he knows he has knowledge the world doesn't want. He's happy someone is "
    "interested in his books, but also very cautious about what he shares."
)

doc.add_heading("Possible Questions", 2)

p = doc.add_paragraph()
p.add_run('"What are dragon riders?"').bold = True
p.add_run(
    "\n\nT'horn will point to the map and tell a short story about ancient dragon riders - they were a group of elves "
    "who bonded themselves to dragons as companions. He will be careful to say this knowledge is forbidden - the "
    "government doesn't like anyone reading about it."
)

p = doc.add_paragraph()
p.add_run('"Why is it forbidden?"').bold = True
p.add_run(
    "\n\nT'horn will be cautious and say he's not completely sure. But he will provide evidence: Over the past "
    "centuries, anyone who researched dragon riders disappeared or went mad. He suspects the government fears a new "
    "rise of dragon riders."
)

p = doc.add_paragraph()
p.add_run('"Do I know someone who knows more?"').bold = True
p.add_run(
    "\n\nT'horn will invite them to leave the city and search for the scholar exiled from the academy. She lives "
    "outside the city, in nature. She can teach them much more."
)

p = doc.add_paragraph()
p.add_run('"What about the map we found?"').bold = True
p.add_run(
    "\n\nT'horn will be SHOCKED - he doesn't know someone was in his library. He will be disturbed but also very "
    "interested. Someone was here... searching for something. He suspects it's connected to dragon riders."
)

# Scene 3
doc.add_heading("Scene 3: Evening on the Road", 1)
doc.add_heading("Setting", 2)
doc.add_paragraph(
    "As the party travels out of the city, they stop to sleep at night. At midnight, they hear sounds nearby - sounds "
    "that seem like small screams (kobolds) and rough sounds (something else, more threatening)."
)

doc.add_heading("The Encounter", 2)
encs = [
    "If they follow the sounds, they arrive at a stone cave",
    "Inside: Kobolds (from the mine) and new, stronger, vicious creatures - Darklings",
    "The kobolds are working on their plan: to capture the scholar outside the city",
    "They say the Betrayer wants information from her - she knows too much about the blood",
    "They're planning to leave this evening or tomorrow"
]
for e in encs:
    doc.add_paragraph(e, style='List Bullet')

doc.add_heading("DM Notes", 2)
dm_notes = [
    "This is not to force combat - it's revelation",
    "The goal is to show the Betrayer is actively searching",
    "This also exposes Darklings as new helpers to the Betrayer",
    "It gives momentum: we need to find the scholar before they do"
]
for note in dm_notes:
    doc.add_paragraph(note, style='List Bullet')

# Bard's Item
doc.add_heading("Bard's Item: Lyre of Harmonic Resonance", 1)
doc.add_heading("Ability", 2)
doc.add_paragraph(
    "When played, it summons spectral musicians (violin, drums, bass, flute) who play in perfect harmony with the Bard. "
    "The music creates a compelling atmosphere and boosts morale."
)

doc.add_heading("Effects", 2)
effects = [
    "Once per long rest, the Bard can activate the lyre (bonus action)",
    "Duration: 5 minutes or until the Bard chooses to stop",
    "Effect: The Bard's Bardic Inspiration dice increase by one size (6d6 → 8d6, 8d8 → 10d8, etc.)",
    "While active, allies within 30 feet get advantage on saves against fear",
    "The music is audible but not loud (doesn't wake neighbors)",
    "The band plays what the Bard wants"
]
for effect in effects:
    doc.add_paragraph(effect, style='List Bullet')

# Session End
doc.add_heading("Session End Goals", 1)
doc.add_paragraph("Players should leave the session with:")
end_goals = [
    "Understanding that dragon riders were real and important",
    "Questions about why they were banned",
    "Knowledge that someone is searching (the map in the library)",
    "Urgency to find the scholar before the kobolds do",
    "Herald: beginning to see a new path (but not fully yet)",
    "Aerendil: more questions about his pact",
    "Bard: more integrated into the group through sharing his story"
]
for goal in end_goals:
    doc.add_paragraph(goal, style='List Bullet')

doc.add_paragraph("\nExpected Duration: 2.5-3 hours")

# Save
output_path = r"c:\עמרי\Shmags 2\projects\dnd-campaign\sessions\session_1_plan.docx"
doc.save(output_path)
print(f"Document created: {output_path}")
