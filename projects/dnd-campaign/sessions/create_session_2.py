from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('Chains', 0)

subtitle = doc.add_heading('Adventure 2 - Session 2', level=2)

# Goals
doc.add_heading('Session Goals', 1)
goals = [
    "Resolve the combat cliffhanger from Session 1",
    "Interrogation reveals: Scholar is being hunted, deadline is tomorrow night, religious orders backed her exile",
    "Players discover the Darkling binding mechanic (crystalline marks)",
    "Aerendil: Dream 1 - one word, wakes next to fire that wasn't there",
    "Herald: learns his order was actively complicit in suppressing Dragon Rider knowledge",
    "Players leave heading toward the Scholar with urgency and a clear deadline",
    "Burton earns his place as a reliable guide"
]
for goal in goals:
    doc.add_paragraph(goal, style='List Bullet')

doc.add_paragraph()

# Scene 1
doc.add_heading('Scene 1: The Cavern Fight', 1)

doc.add_heading('Setting', 2)
doc.add_paragraph(
    "The cavern mouth. Darklings emerged from the dark after Aerendil's yell, short swords drawn. "
    "Kobolds cluster behind them - more scared than aggressive."
)

doc.add_heading('What Happens', 2)
s1_events = [
    "Darklings are the real threat - disciplined, dangerous, not easily scared",
    "Kobolds fold early - one breaks from the group in the first or second round and drops its weapon",
    "When players get within a few feet of a Darkling, they notice: faint crystalline marks on the skin, almost like veins of pale blue",
    "Combat ends with at least one captive (the Kobold that broke, or a Darkling if players subdue one)"
]
for e in s1_events:
    doc.add_paragraph(e, style='List Bullet')

doc.add_heading('DM Notes', 2)
s1_notes = [
    "Keep it fast: 2-3 meaningful rounds, not a war of attrition",
    "The crystalline marks are visible but say nothing about them - let players ask or wonder",
    "If all enemies are killed with no captive: one Darkling's body has a crude map with a location marked east in the mountains",
    "The Kobold that broke is genuinely terrified - it will cooperate fast once it believes it can survive"
]
for note in s1_notes:
    doc.add_paragraph(note, style='List Bullet')

doc.add_paragraph()

# Scene 2
doc.add_heading('Scene 2: Interrogation', 1)

doc.add_heading('Setting', 2)
doc.add_paragraph(
    "Outside the cavern or just inside, captive restrained. Night, cold mountain air."
)

doc.add_heading('The Kobold (breaks easily)', 2)
doc.add_paragraph(
    "Offer it its life in exchange for information. It gives everything without much pressure:"
)
kobold_lines = [
    '"She lives east - in the mountains, away from the city. She came from the academy."',
    '"He wants her alive. She knows the ritual. She knows how the blood works, how to make the bond."',
    '"We were supposed to move on her tonight. There are more of us - the main group moves tomorrow night."',
    '"The orders in the city signed off on her exile. The academy didn\'t act alone."'
]
for line in kobold_lines:
    doc.add_paragraph(line, style='List Bullet')

doc.add_heading('The Darkling (one piece, under real pressure only)', 2)
doc.add_paragraph(
    "The Darkling won't cooperate willingly. Under serious pressure it gives one thing:"
)
doc.add_paragraph('"There are more. The main group moves tomorrow night."', style='List Bullet')
doc.add_paragraph(
    "Then stops. Not out of loyalty - the binding costs it when it tries to say more. "
    "Players may notice the crystalline marks pulse faintly when it resists. "
    "The Darkling won't give more. It isn't braver than the Kobold - it physically can't. "
    "That wrongness should feel like information in itself."
)

doc.add_heading('Character Moment - Herald', 2)
p = doc.add_paragraph()
p.add_run('"The orders in the city signed off on her exile."').bold = True
p.add_run(
    "\n\nHe doesn't need to recognize anything specific. He just learns, for the first time, "
    "that the institution he came from was active in this suppression. Not passive. Not ignorant. "
    "They signed off. Let him sit with it. Don't rush past it."
)

doc.add_heading('DM Notes', 2)
s2_notes = [
    "Kobold gives everything freely once scared enough",
    "Darkling gives one piece of information maximum",
    "The binding mechanic should feel wrong, not explained",
    "Don't underline the Herald moment - let his player react naturally"
]
for note in s2_notes:
    doc.add_paragraph(note, style='List Bullet')

doc.add_paragraph()

# Scene 3
doc.add_heading('Scene 3: Rest in the Mountains', 1)

doc.add_heading('Setting', 2)
doc.add_paragraph(
    "The group finds a place to rest away from the cavern. Cold mountain air, quiet, open sky."
)

doc.add_heading('Burton', 2)
doc.add_paragraph(
    "When asked about the Scholar's general area, Burton mentions almost offhandedly that he's seen "
    "smoke from a camp east that doesn't match any trade route. He assumed it was a hermit. "
    "He can take them there."
)

doc.add_heading("Aerendil's Dream 1", 2)
doc.add_paragraph("During watch or early sleep, the dream hits.")
doc.add_paragraph()

dream_p = doc.add_paragraph()
dream_run = dream_p.add_run(
    "Heat. Not painful - the opposite. A warmth that comes from somewhere far below everything, "
    "pressing upward. A presence at the edge of it, enormous and old and very far away. "
    "Then one word, not quite heard but felt: "
)
dream_run.italic = True
chains_run = dream_p.add_run('"Chains."')
chains_run.italic = True
chains_run.bold = True
end_run = dream_p.add_run(
    " He wakes up. There is a small fire beside him. There was no fire before."
)
end_run.italic = True

doc.add_paragraph()
dream_notes = [
    "No option to continue yet - this is Dream 1, just the opening",
    "He knows it was real. He knows it connects to the crystal and the fire from the night before.",
    "The fire is small, warm, non-threatening. No explanation. It just is."
]
for dn in dream_notes:
    doc.add_paragraph(dn, style='List Bullet')

doc.add_heading('DM Notes', 2)
s3_notes = [
    "Don't rush this scene - it's where character moments breathe",
    "Let players lead the conversation after the interrogation",
    "The deadline should stay present: they need to move before tomorrow night"
]
for note in s3_notes:
    doc.add_paragraph(note, style='List Bullet')

doc.add_paragraph()

# Scene 4
doc.add_heading('Scene 4: The Journey East', 1)

doc.add_heading('Setting', 2)
doc.add_paragraph(
    "Mountain terrain. Burton leads east through paths no one else would know. "
    "Steep, rocky, quiet except for wind."
)

doc.add_heading('What Happens', 2)
s4_events = [
    "Burton navigates confidently - this is his terrain. Players start to trust him practically even if they still find him odd.",
    "At some point they pass ruins of an old shrine. Predates the current order by centuries. The imagery shows humans alongside creatures - not commanding them, alongside them. Partners. Nobody explains it. Herald sees it.",
    "End the session arriving within sight of a small camp with smoke rising (the Scholar's), OR finding fresh Darkling signs close to where she is - recent tracks, a dropped weapon, disturbed ground."
]
for e in s4_events:
    doc.add_paragraph(e, style='List Bullet')

doc.add_heading('DM Notes', 2)
s4_notes = [
    "The shrine is a visual beat, not a lecture. No explanation - Herald's player will understand.",
    "Keep Burton's dialogue practical and slightly odd - useful but still strange",
    "The deadline matters: Burton can reference it ('we need to move if we want to arrive before they do')",
    "Either session ending creates the same hook: she's real, she's close, and they're not the only ones looking"
]
for note in s4_notes:
    doc.add_paragraph(note, style='List Bullet')

doc.add_paragraph()

# Reference: Darkling Binding
doc.add_heading('Reference: The Darkling Binding', 1)
doc.add_paragraph(
    "The Betrayer has been using blue salt crystal magic to bind the Darklings to him - a forced, "
    "non-consensual magical connection. The same principle he intends to apply to the ancient dragon, "
    "tested first on smaller creatures."
)

p = doc.add_paragraph()
p.add_run('Visual: ').bold = True
p.add_run('Faint crystalline marks on the Darkling\'s skin, like pale blue veins. More visible when it tries to resist.')

p = doc.add_paragraph()
p.add_run('Effect: ').bold = True
p.add_run('Darklings cannot betray the Betrayer without the binding physically resisting. The cost increases the more they fight it.')

p = doc.add_paragraph()
p.add_run('What Players Know Right Now: ').bold = True
p.add_run('Something is wrong with the Darklings. The marks are there. Explanation comes later.')

p = doc.add_paragraph()
p.add_run('Thematic Meaning: ').bold = True
p.add_run("The Darklings aren't just loyal soldiers - they're victims of the same magic threatening the dragon.")

doc.add_paragraph()

# Reference: Dream System
doc.add_heading("Reference: Aerendil's Dream System", 1)

dream_rows = [
    ("Dream 1 (this session)", "Heat, presence, one word (\"chains\"), wakes next to fire that wasn't there. No option to continue."),
    ("Dream 2 onward", "Option to continue introduced. Each time he continues, wakes up closer to the fire."),
    ("Final dream", "Goes fully into the fire. The patron reveals itself."),
    ("Clarity progression", "Sensations → single words → images → fragments → direct contact. Dreams get clearer as campaign progresses.")
]

table = doc.add_table(rows=len(dream_rows) + 1, cols=2)
table.style = 'Light Grid Accent 1'

hdr = table.rows[0].cells
hdr[0].text = 'Dream'
hdr[1].text = 'Description'

for i, (dream, desc) in enumerate(dream_rows, 1):
    cells = table.rows[i].cells
    cells[0].text = dream
    cells[1].text = desc

doc.add_paragraph()

# Session End
doc.add_heading('Session End Goals', 1)
doc.add_paragraph('Players should leave the session knowing:')
end_goals = [
    "The Scholar is real, east in the mountains, and in immediate danger",
    "Tomorrow night is the deadline - the main Darkling group moves then",
    "Something is wrong with the Darklings (the marks)",
    "Aerendil felt something real in the dream - and woke next to a fire that wasn't there",
    "Herald's order wasn't passive - they signed off on the Scholar's exile",
    "Burton knows this terrain and he's useful"
]
for goal in end_goals:
    doc.add_paragraph(goal, style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Expected Duration: 2.5 hours').bold = True

output_path = r"c:\עמרי\Shmags 2\projects\dnd-campaign\sessions\session_2_plan.docx"
doc.save(output_path)
print(f"Document created: {output_path}")
