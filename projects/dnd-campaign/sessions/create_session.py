#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('סשן ספריית אחוזת תורן', 0)
title.alignment = WD_ALIGN_PARAGRAPH.RIGHT

subtitle = doc.add_heading('הרפתקה 1: כרך 2 - דרקון רוכבים הקדומים', level=2)
subtitle.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Goals
goals_h = doc.add_heading('מטרות הסשן:', level=1)
goals_h.alignment = WD_ALIGN_PARAGRAPH.RIGHT

goals = [
    "הכרת תרבות ודברים היסטוריים על דרקון רוכבים",
    "חשיפת תפקיד גבישי המלח הכחולים בריכוך דרקונים",
    "שיפור התואר ההתחלתי של Herald בדרך חדשה",
    "התחלת קשר בין Aerendil ובעל חוזה הוורלוק",
    "חשיפת סיפור העבר של הבארד ודילו עם תורן",
    "גילוי מפה המראה עקבות של חוקר בספריה",
    "התחזוקה על דרקון רוכבים כנושא מרכזי בקמפיין",
    "הצגת המלומדת הגלויה בחוץ לעיר כאפשרות הדרכה"
]

for goal in goals:
    p = doc.add_paragraph(goal, style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.add_paragraph()

# Scene 1: Library
scene1 = doc.add_heading('סצנה 1: ספריית אחוזת תורן', level=1)
scene1.alignment = WD_ALIGN_PARAGRAPH.RIGHT

desc_h = doc.add_heading('הגדרה:', level=2)
desc_h.alignment = WD_ALIGN_PARAGRAPH.RIGHT

description = doc.add_paragraph('הספריה של תורן היא חלל ענק, עם מדפים גבוהים המגיעים לתקרה, אור טבעי זורם דרך חלונות גדולים עם וילונות כבדים. על הרצפה יש ערימות של ספרים, מפות, ושלדים של פריטים קסומים מכוסים באבק. ריח של נייר ישן וגבס מלא את האוויר. הספריה היא מערך קטוגוס, כמעט כאוטית - אבל לאדם שיודע לחפש, יש אוצר של ידע.')
description.alignment = WD_ALIGN_PARAGRAPH.RIGHT

opp_h = doc.add_heading('הזדמנויות:', level=2)
opp_h.alignment = WD_ALIGN_PARAGRAPH.RIGHT

opportunities = [
    "שלוש סיפורים היסטוריים על דרקון רוכבים (ספר עתיק, ציורים, כתיבת יד)",
    "גבישים כחולים קטנים בקופסה מזכוכית (לא ברור למה הם כאן)",
    "פריטים קסומים נמוכים (אם יוציאו DC 16 או 20+)",
    "דברים מונדיים - ספרים, מפות, פחיות, מטבעות, שרדים",
    "הזדמנות לבארד לדבר על עברו אם הוא רוצה",
    "הודעה של תורן שהוא מעט בורח משימוש הגוארד (מוצא לנושא)"
]

for opp in opportunities:
    p = doc.add_paragraph(opp, style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

char_h = doc.add_heading('אלמנטים ספציפיים לדמויות:', level=2)
char_h.alignment = WD_ALIGN_PARAGRAPH.RIGHT

herald_p = doc.add_paragraph()
herald_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
herald_run = herald_p.add_run('Herald: ')
herald_run.bold = True
herald_p.add_run('צריך להיות נחשף לסיפור ההרואי של דרקון רוכבים. צריך לראות את ההבדל בין הערכים של דרקון רוכבים (חופש, שותפות) לבין הסדר שעזב. זה צריך להתחיל לתעורר קול פנימי של "זה משהו שאני יכול להיות"')

aerendil_p = doc.add_paragraph()
aerendil_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
aerendil_run = aerendil_p.add_run('Aerendil: ')
aerendil_run.bold = True
aerendil_p.add_run('צריך לגלות את הסיפור על תורשת האש של דרקון רוכבים. צריך להבין שהישות המסתורית שלו עשויה להיות דרקון עתיק. צריך לתרגם את החלומות האדומים/האש כחלק מהחוזה הזה. לא צריך להבין הכל עדיין - יש עדיין מסתורין.')

bard_p = doc.add_paragraph()
bard_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
bard_run = bard_p.add_run('הבארד: ')
bard_run.bold = True
bard_p.add_run('השתקט יחסית בזמן הספריה. בשלב מסוים (אחרי שתורן חושף את הדיל), הוא יכול לשתף על עברו וההרכב הקודם שלו. זה צריך להיות רגע קבוצתי.')

notes_h = doc.add_heading('הערות למנחה:', level=2)
notes_h.alignment = WD_ALIGN_PARAGRAPH.RIGHT

notes = [
    "אם הם רוצים לדבר אם משהו, תן להם מרחק לגלות",
    "אם הם נראים איבודים, תן להם רמזים קטנים",
    "המפה צריכה להישמר עבור סוף הסשן",
    "כשהם מוצאים את הגבישים הכחולים - זה רמז קטנה שהם קשורים לדרקון רוכבים"
]

for note in notes:
    pn = doc.add_paragraph(note, style='List Bullet')
    pn.alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.add_paragraph()

# Scene 2: T'horn
scene2 = doc.add_heading('סצנה 2: מפגש עם תורן', level=1)
scene2.alignment = WD_ALIGN_PARAGRAPH.RIGHT

thorn_h = doc.add_heading('הגדרה:', level=2)
thorn_h.alignment = WD_ALIGN_PARAGRAPH.RIGHT

thorn_desc = doc.add_paragraph('תורן הוא גבר בגיל 50-60, מלא בחוזק אבל מעט מבולגן בהופעתו. הוא אוהב להשמע חושב אך גם עצוב - יודע שיש בו ידע שהעולם אינו רוצה. הוא שמח שיש מישהו שעניין בספרים שלו, אך גם זהיר מאוד לגבי מה הוא משתף.')
thorn_desc.alignment = WD_ALIGN_PARAGRAPH.RIGHT

q_h = doc.add_heading('שאלות אפשריות:', level=2)
q_h.alignment = WD_ALIGN_PARAGRAPH.RIGHT

questions = [
    ("מה זה דרקון רוכבים?", "תורן יצביע למפה והוא יספר סיפור קצר על דרקון רוכבים העתיקים - הם היו קבוצה של אלפים המשמשים כגברים שחיברו את עצמם לדרקונים. הוא יהיה זהיר להגיד שהידע הזה אסור - הממשלה לא אוהבת מישהו שקורא על זה."),
    ("מדוע זה אסור?", "תורן יהיה זהיר ויגיד שהוא לא בטוח לחלוטין. אבל הוא יספק עדות: במהלך המאות השנים האחרונות, כל מי שנחקר על ידי דרקון רוכבים נעלם או נזקק. הוא יחשוד שהממשלה חוששת מעלייה חדשה של דרקון רוכבים."),
    ("אני מכיר מישהו שיודע יותר?", "תורן יזמין אותם שיעזבו את העיר ויחפשו את המלומדת שגורשה מהאקדמיה. היא חיה בחוץ לעיר, בטבע. היא יכולה ללמד אותם הרבה יותר."),
    ("מה עם המפה שמצאנו?", "תורן יהיה SHOCKED. הוא לא יודע שמישהו היה בספריה שלו. הוא יהיה מוטרד אך גם מעניין - מישהו היה כאן... מחפש משהו. הוא יחשוד שזה קשור לדרקון רוכבים.")
]

for q, a in questions:
    pq = doc.add_paragraph()
    pq.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rq = pq.add_run(f'"{q}"')
    rq.bold = True

    pa = doc.add_paragraph(a, style='List Bullet')
    pa.alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.add_paragraph()

# Scene 3: Road Encounter
scene3 = doc.add_heading('סצנה 3: הערב בדרך', level=1)
scene3.alignment = WD_ALIGN_PARAGRAPH.RIGHT

road_h = doc.add_heading('הגדרה:', level=2)
road_h.alignment = WD_ALIGN_PARAGRAPH.RIGHT

road_desc = doc.add_paragraph('כשהשחקנים חוברים בדרך מהעיר, הם עוצרים לשינה בלילה. בשעת חצי לילה, הם שומעים קולות מהקרוב - הנשמע כמו קולות של צעקות קטנות (קובולדים) וקולות מחוספסים (משהו אחר, יותר מאיים).')
road_desc.alignment = WD_ALIGN_PARAGRAPH.RIGHT

disc_h = doc.add_heading('ההתגלות:', level=2)
disc_h.alignment = WD_ALIGN_PARAGRAPH.RIGHT

discovery = [
    "אם הם עוקבים אחרי הקולות, הם מגיעים לערת סלע",
    "בתוך הערה: קובולדים (מהמכרה) וקובולדים חדשים - Darklings - יצורים שחורים חזקים וזדוניים יותר",
    "הקובולדים בתחזוקה אחרת דברים על תוכניתם: לתפוס את המלומדת בחוץ לעיר",
    "הם אומרים שהבוגדן רוצה מידע ממנה - היא יודעת הרבה מדי על הדם",
    "הם תוכננים ללכת לערב הזה או מחר"
]

for item in discovery:
    pd = doc.add_paragraph(item, style='List Bullet')
    pd.alignment = WD_ALIGN_PARAGRAPH.RIGHT

road_notes_h = doc.add_heading('הערות:', level=2)
road_notes_h.alignment = WD_ALIGN_PARAGRAPH.RIGHT

road_notes = [
    "זה לא כדי להכריח קרב - זה גילוי",
    "המטרה היא להראות שהבוגדן מחפש בפעילות",
    "זה גם חושף Darklings כעוזרים חדשים לבוגדן",
    "זה נותן דחיפה: אנחנו צריכים למצוא את המלומדת לפני שהם יעשו זאת"
]

for rn in road_notes:
    prn = doc.add_paragraph(rn, style='List Bullet')
    prn.alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.add_paragraph()

# Magic Items
mag_h = doc.add_heading('פריטים קסומים נמוכים', level=1)
mag_h.alignment = WD_ALIGN_PARAGRAPH.RIGHT

mag_note = doc.add_paragraph('אם יוציאו DC 16 בחקירה: 1 פריט. אם DC 20+: 3 פריטים')
mag_note.alignment = WD_ALIGN_PARAGRAPH.RIGHT

table1 = doc.add_table(rows=4, cols=2)
table1.style = 'Light Grid Accent 1'

hc = table1.rows[0].cells
hc[0].text = 'פריט'
hc[1].text = 'תיאור'

for cell in hc:
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

items = [
    ('אמולט של סיום', 'למטה יש כושר עדיפות על שמירות נגד רעל'),
    ('טבעת חום', 'למטה עמיד בנזקי קור לא קסום'),
    ('גלימת אלפינקינד', 'למטה יש כושר עדיפות בבדיקות Stealth בעת אור טבעי')
]

for i, (item, desc) in enumerate(items, 1):
    cells = table1.rows[i].cells
    cells[0].text = item
    cells[1].text = desc
    cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.add_paragraph()

# Mundane Items
mund_h = doc.add_heading('פריטים מונדיים בספריה', level=1)
mund_h.alignment = WD_ALIGN_PARAGRAPH.RIGHT

table2 = doc.add_table(rows=11, cols=2)
table2.style = 'Light Grid Accent 1'

hm = table2.rows[0].cells
hm[0].text = 'מס'
hm[1].text = 'פריט'

for cell in hm:
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

mundane_items = [
    'ספר סחר נתונים - רשומות של נתיבי סוחרים דרך הממלכות',
    'לוחות אבן - חוקים עתיקים של Clerroc (שחוקים, קשה לקרוא)',
    'שרדי כלים - שברי קרמיקה מתקופת ייסוד Clerroc',
    'תוכניות אדריכליות - עיצובי מערכת מים של Clerroc (מפורטים וטכניים)',
    'יומן של סוחר - ערכים אישיים מ-200 שנה אחרונות',
    'אוסף של מטבעות עתיקים - מטבע מתקופות ואזורים שונים',
    'חיבור צבאי - כתיבות אסטרטגיות על מצור קדום',
    'מדריך אלכימיה - מדריך מאויר על צמחים וריפויים מקומיים',
    'מכתבים בין בית אצילות - התכתבות על סכסוכי סחר',
    'אוסף מפות של קרטוגרף - מפות שונות של האזור (חלק מהן מיושנות)'
]

for i, item in enumerate(mundane_items, 1):
    cells = table2.rows[i].cells
    cells[0].text = str(i)
    cells[1].text = item
    cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.add_paragraph()

# Bard Instrument
bard_h = doc.add_heading('כלי הנגינה של הבארד: הליר ההרמוניה של תהודה', level=1)
bard_h.alignment = WD_ALIGN_PARAGRAPH.RIGHT

ability_h = doc.add_heading('כושר:', level=2)
ability_h.alignment = WD_ALIGN_PARAGRAPH.RIGHT

ability = doc.add_paragraph('כשמנוגן, קורא למוזיקאים ספקטרלים (כינור, תופים, בס, חליל) המנגנים בהרמוניה מושלמת עם הבארד. המוזיקה יוצרת אווירה משכנעת וממגבר מורל.')
ability.alignment = WD_ALIGN_PARAGRAPH.RIGHT

mech_h = doc.add_heading('מכניקה (D&D 5e):', level=2)
mech_h.alignment = WD_ALIGN_PARAGRAPH.RIGHT

mechanics = [
    "פעם לאחר שיום ארוך, הבארד יכול להפעיל את הליר (פעולת בונוס)",
    "משך זמן: 5 דקות או עד שהבארד בוחר להרוג אותה",
    "השפעה: גבישי Bardic Inspiration של הבארד גדלים בגודל אחד (d6 → d8, d8 → d10, וכו')",
    "כל בעוד הוא פעיל, בעלי אחות 30 רגל מקבלים כושר עדיפות על שמירות נגד פחד",
    "המוזיקה נשמעת אבל לא חזקה (לא מעיר משכנים)",
    "הלהקה נוגנת מה שהבארד רוצה"
]

for mech in mechanics:
    pm = doc.add_paragraph(mech, style='List Bullet')
    pm.alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.add_paragraph()

# Conclusion
conclusion_h = doc.add_heading('סיום הסשן:', level=1)
conclusion_h.alignment = WD_ALIGN_PARAGRAPH.RIGHT

conclusion = doc.add_paragraph('השחקנים צריכים להשאיר את הסשן עם: הבנה שדרקון רוכבים היו אמתיים וחשובים, משאלת דעות על מדוע הם נאסרו, ידע שמישהו מחפש (המפה בספריה), דחיפה לחפש את המלומדת לפני שהקובולדים וDarklings יעשו זאת, Herald מתחיל לראות דרך חדשה (אבל לא מלא עדיין), Aerendil עם יותר שאלות על הוורלוק שלו, והבארד יותר משולב בקבוצה דרך שיתוף הסיפור שלו.')
conclusion.alignment = WD_ALIGN_PARAGRAPH.RIGHT

duration = doc.add_paragraph('\nמשך זמן צפוי: 2.5-3 שעות')
duration.alignment = WD_ALIGN_PARAGRAPH.RIGHT
duration.runs[0].bold = True

doc.save('c:\\עמרי\\Shmags 2\\session_plan.docx')
print("Document created successfully!")
