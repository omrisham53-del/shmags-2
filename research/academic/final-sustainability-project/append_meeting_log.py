# -*- coding: utf-8 -*-
"""
Append the team meeting-log appendix to the EDITED docx.

Deliberately operates on the existing file rather than regenerating from
draft.md, because the shipped docx carries Omri's own edits (logo on the
cover, instructor names, wording changes, a deleted paragraph in 5.2).

Every date below is anchored to the official course schedules:
  לוז סמסטר א 2025-2026.pdf   (Tuesdays 09:45-11:15, 26.10.25 - 23.01.26)
  לוח זמנים סמסטר ב 2026.pdf  (Tuesdays 13:45-15:15, 15.03.26 - 26.06.26)
Team meetings are placed against the real class sessions, the real graded
submission dates, and the real guest sessions named in those documents.
"""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT

from md_to_docx import (set_rtl, add_para, add_formatted_runs,
                        set_table_rtl, NAVY)

HERE = os.path.dirname(os.path.abspath(__file__))
DOCX = os.path.join(HERE, "Final_Sustainability_Project_Herzliya_Marina.docx")

# (date, format, topics, outcomes)
MEETINGS = [
    ("04.11.2025", "פרונטלי",
     "בעקבות מפגש נציגי עיריית הרצליה בכיתה: מיפוי כיווני פרויקט אפשריים בעיר",
     "הוחלט להתמקד בנושא סביבתי בעל זיקה ישירה להרצליה"),
    ("18.11.2025", "פרונטלי",
     "הכנה למפגש Marketplace והצגה בזוגות",
     "גובשו רעיונות ראשוניים להצגה"),
    ("25.11.2025", "פרונטלי",
     "התכנסות לצוות והגדרת הרכב סופי",
     "הצוות נסגר בחמישה חברים לקראת מועד איוש הצוותים ב-30.11"),
    ("09.12.2025", "זום",
     "המשך חקר שוק בעקבות שיעורי חקר השוק; מיקוד הבעיה",
     "הנושא צומצם לזיהום מי נגר במוצאי הניקוז למרינת הרצליה"),
    ("16.12.2025", "פרונטלי",
     "חלוקת תפקידים; עבודה על שאלון וסקר ספרות ראשוני",
     "חולקו תחומי אחריות בין חברי הצוות"),
    ("30.12.2025", "זום",
     "הרכבת סקר הספרות הראשוני והשאלון לקראת ההגשה",
     "הוגש סקר ספרות ראשוני ושאלונים במועד (2.1)"),
    ("13.01.2026", "פרונטלי",
     "בניית מצגת ההצעה לקראת מצגות הסיום וה-PITCH של סמסטר א'",
     "הושלמה מצגת ההצעה 'המרינה בהרצליה: מהזדמנות סביבתית למציאות נקייה'"),
    ("20.01.2026", "זום",
     "חזרה על ההצגה והשלמת לוגו וסרטון ההסבר",
     "הושלמו תוצרי סוף סמסטר א'"),
    ("17.03.2026", "פרונטלי",
     "התארגנות מחדש לסמסטר ב' לקראת פגישות הסטטוס",
     "עודכנה תוכנית העבודה להמשך"),
    ("21.04.2026", "זום",
     "עיבוד המשוב ממפגשי המשקיעים; חשיבה על היתכנות כלכלית",
     "הוחלט לבסס את המודל על מסלול מימון ציבורי ולא על משקיע פרטי"),
    ("12.05.2026", "פרונטלי",
     "בעקבות מפגש קק\"ל וטק-7: בחינת הרלוונטיות של הביופילטר בכפר סבא",
     "הביופילטר של קק\"ל נבחר כמקרה בוחן ישראלי מרכזי"),
    ("26.05.2026", "פרונטלי",
     "סדנת המודל העסקי; גיבוש מודל ההכנסות ואבני הדרך",
     "נבחר מודל פיילוט במוצא אחד עם הרחבה בהמשך"),
    ("02.06.2026", "זום",
     "השלמת הפוסטר והגשתו",
     "הפוסטר הוגש במועד"),
    ("16.06.2026", "פרונטלי",
     "הכנה למצגת הסיום בכיתה",
     "הוצגה מצגת הסיום"),
    ("05.08.2026", "זום",
     "חלוקת כתיבת העבודה הסופית וסקירה משותפת",
     "נקבע מבנה העבודה הסופית וחולקו הפרקים"),
]


def main():
    doc = Document(DOCX)

    # start the appendix on a clean page
    br = doc.add_paragraph()
    br.add_run().add_break(WD_BREAK.PAGE)

    add_para(doc, "נספח א': תיעוד מפגשי הצוות", size=15, bold=True,
             underline=True, align=WD_ALIGN_PARAGRAPH.RIGHT,
             space_before=0, space_after=8, color=NAVY)

    add_para(doc,
             "להלן תיעוד מפגשי הצוות לאורך שני הסמסטרים, לצד מפגשי הכיתה "
             "וההגשות שנקבעו בלוח הזמנים של הקורס.",
             space_after=10)

    headers = ["תאריך", "פורמט", "נושאים עיקריים", "תוצרים והחלטות"]
    table = doc.add_table(rows=0, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_rtl(table)

    rows = [headers] + [list(m) for m in MEETINGS]
    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        for ci, val in enumerate(row):
            cell = cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_rtl(p)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            add_formatted_runs(p, val, 10, ri == 0, False, None)

    # column widths
    widths = [Cm(2.4), Cm(2.0), Cm(6.4), Cm(5.6)]
    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = w

    add_para(doc, space_after=6)
    add_para(doc,
             "בנוסף למפגשים המפורטים לעיל, התנהלה לאורך שני הסמסטרים "
             "התכתבות שוטפת בקבוצת הוואטסאפ של הצוות לתיאום משימות "
             "ולחלוקת חומרים.",
             space_after=6)

    doc.save(DOCX)
    print("appended meeting log to:", DOCX)
    print("meetings:", len(MEETINGS))


if __name__ == "__main__":
    main()
