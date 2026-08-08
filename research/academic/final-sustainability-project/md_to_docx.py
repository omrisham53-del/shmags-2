# -*- coding: utf-8 -*-
"""
Build the submission-ready Hebrew RTL .docx from draft.md.

Brief's technical requirements: 1.15 line spacing, font 12, up to 20 pages
excluding appendices and bibliography.

Style guide (references/academic-style-guide.md, Hebrew Mode):
  - section headers bold + underlined, right-aligned
  - body text justified
  - RTL paragraph direction throughout; LTR only where explicitly needed
"""
import os
import re
import io

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(HERE, "draft.md")
CHARTS = os.path.join(HERE, "charts")
OUT = os.path.join(HERE, "Final_Sustainability_Project_Herzliya_Marina.docx")

BODY_FONT = "David"        # clean Hebrew serif, present on this machine
BODY_SIZE = 12
NAVY = RGBColor(0x12, 0x3B, 0x5C)

# charts are injected after the paragraph whose text starts with these keys
CHART_ANCHORS = [
    ("03_outlet_matrix.png",
     "בנוסף, נוהל ניקיון ושאיבת פסולת מהנקזים",
     "איור 1: מצב האמצעים בארבעת מוצאי הניקוז אל הים"),
    ("01_bathing_season_exceedance.png",
     "זהו הממצא המרכזי שעליו נשענת",
     "איור 2: שיעור הדגימות החורגות בחופי הרצליה בעונת הרחצה"),
    ("02_removal_efficiency.png",
     "התוצאות היו עקביות לאורך תקופת הניטור",
     "איור 3: יעילות הרחקת מזהמים במערכת מדיה פילטר"),
    ("04_treatment_train.png",
     "**שלב 4: תפעול ותחזוקה.**",
     "איור 4: רכבת הטיפול המוצעת במוצא הניקוז"),
]


# --------------------------------------------------------------- rtl helpers
def set_rtl(paragraph):
    """Mark a paragraph as right-to-left at the XML level."""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)


def style_run(run, size=BODY_SIZE, bold=False, underline=False,
              color=None, font=BODY_FONT):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.underline = underline
    run.font.name = font
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    # Hebrew text lives in the "complex script" slot; set both.
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:cs"), font)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(int(size * 2)))
    rPr.append(szCs)
    if bold:
        bCs = OxmlElement("w:bCs")
        rPr.append(bCs)
    rtl = OxmlElement("w:rtl")
    rtl.set(qn("w:val"), "1")
    rPr.append(rtl)


def add_para(doc, text="", size=BODY_SIZE, bold=False, underline=False,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6,
             space_before=0, color=None, spacing=1.15):
    p = doc.add_paragraph()
    p.alignment = align
    set_rtl(p)
    pf = p.paragraph_format
    pf.line_spacing = spacing
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    if text:
        add_formatted_runs(p, text, size, bold, underline, color)
    return p


def add_formatted_runs(p, text, size, bold, underline, color):
    """Handle inline **bold** segments inside a line of markdown."""
    for seg in re.split(r"(\*\*[^*]+\*\*)", text):
        if not seg:
            continue
        if seg.startswith("**") and seg.endswith("**"):
            r = p.add_run(seg[2:-2])
            style_run(r, size, True, underline, color)
        else:
            r = p.add_run(seg)
            style_run(r, size, bold, underline, color)


def set_table_rtl(table):
    tblPr = table._tbl.tblPr
    bidi = OxmlElement("w:bidiVisual")
    tblPr.append(bidi)


# --------------------------------------------------------------- cover page
def build_cover(doc):
    for _ in range(4):
        add_para(doc, space_after=0)

    add_para(doc, "מים צלולים", size=30, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6, color=NAVY)
    add_para(doc, "טיהור מי נגר עירוני במוצאי הניקוז למרינת הרצליה",
             size=17, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=42)

    add_para(doc, "פרויקט יישומי בקיימות", size=14, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=36)

    for name, idnum in [
        ("תום גינל", "208747881"),
        ("תומר טסה", "209164458"),
        ("אבישי מאייר", "209498039"),
        ("עמרי שמגר", "323081851"),
        ("יונתן חורי", "323915546"),
    ]:
        add_para(doc, f"{name}  {idnum}", size=12,
                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)

    for _ in range(3):
        add_para(doc, space_after=0)

    add_para(doc, "אוניברסיטת רייכמן", size=12,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    add_para(doc, "מרצה: [שם המרצה]", size=12,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    add_para(doc, "תאריך הגשה: 13/08/2026", size=12,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


# --------------------------------------------------------------- md parsing
def flush_table(doc, rows):
    """rows: list of list-of-cell-strings; first row is the header."""
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_rtl(table)

    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        for ci in range(ncols):
            txt = row[ci] if ci < len(row) else ""
            cell = cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_rtl(p)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            add_formatted_runs(p, txt, 10.5, ri == 0, False, None)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def insert_chart(doc, filename, caption):
    path = os.path.join(CHARTS, filename)
    if not os.path.exists(path):
        print("  ! missing chart:", filename)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(path, width=Cm(15.2))

    cap = add_para(doc, caption, size=10, bold=False,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    for r in cap.runs:
        r.font.italic = True


def build(doc):
    md = io.open(SRC, encoding="utf-8").read()
    lines = md.split("\n")

    # skip the markdown title block; the cover page replaces it
    start = next(i for i, l in enumerate(lines) if l.strip() == "## תקציר מנהלים")
    lines = lines[start:]

    pending_table = []
    used_charts = set()

    def flush():
        nonlocal pending_table
        if pending_table:
            flush_table(doc, pending_table)
            pending_table = []

    for raw in lines:
        line = raw.rstrip()
        stripped = line.strip()

        # tables
        if stripped.startswith("|"):
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            if all(re.fullmatch(r":?-{2,}:?", c) for c in cells if c):
                continue  # separator row
            pending_table.append(cells)
            continue
        flush()

        if not stripped:
            continue

        if stripped.startswith("---"):
            continue

        # italic-only note lines
        if stripped.startswith("*(") and stripped.endswith(")*"):
            p = add_para(doc, stripped.strip("*"), size=10.5,
                         align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=8)
            for r in p.runs:
                r.font.italic = True
            continue

        # headers
        if stripped.startswith("### "):
            add_para(doc, stripped[4:], size=12.5, bold=True,
                     align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=10,
                     space_after=4, color=NAVY)
            continue
        if stripped.startswith("## "):
            title = stripped[3:]
            add_para(doc, title, size=15, bold=True, underline=True,
                     align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=16,
                     space_after=8, color=NAVY)
            continue

        # bullets
        if stripped.startswith("- "):
            p = add_para(doc, stripped[2:], align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                         space_after=3)
            p.paragraph_format.right_indent = Cm(0.75)
            continue
        if re.match(r"^\d+\.\s", stripped):
            p = add_para(doc, stripped, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                         space_after=3)
            p.paragraph_format.right_indent = Cm(0.75)
            continue

        # ordinary paragraph
        add_para(doc, stripped, space_after=7)

        # chart injection
        for fname, anchor, caption in CHART_ANCHORS:
            if fname in used_charts:
                continue
            if stripped.startswith(anchor):
                insert_chart(doc, fname, caption)
                used_charts.add(fname)

    flush()

    missing = [c[0] for c in CHART_ANCHORS if c[0] not in used_charts]
    if missing:
        print("  ! charts never anchored:", missing)


def main():
    doc = Document()

    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Cm(2.0)
    sec.top_margin = sec.bottom_margin = Cm(2.0)
    # right-to-left page direction
    sectPr = sec._sectPr
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    sectPr.append(bidi)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(BODY_SIZE)

    build_cover(doc)
    build(doc)

    doc.save(OUT)
    print("saved:", OUT)


if __name__ == "__main__":
    main()
