# -*- coding: utf-8 -*-
"""Convert draft.md (this assignment's specific structure) into a .docx file."""
import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = r"c:\עמרי\Shmags 2\research\academic\final-lca-assignment"
SRC = os.path.join(BASE, "draft.md")
OUT = os.path.join(BASE, "Final_Assignment_LCA_Comparative_EPD.docx")

INK = RGBColor(0x0B, 0x0B, 0x0B)

with open(SRC, "r", encoding="utf-8") as f:
    text = f.read()

lines = text.split("\n")

doc = Document()

# base style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = INK

for h_name, size in [("Heading 1", 16), ("Heading 2", 13.5), ("Heading 3", 12)]:
    st = doc.styles[h_name]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = INK
    st.font.italic = False


def add_runs_with_italics(paragraph, raw_text, size=None, color=None):
    """Split on *...* markdown italics and add alternating runs."""
    parts = re.split(r"(\*[^*]+\*)", raw_text)
    for part in parts:
        if not part:
            continue
        italic = part.startswith("*") and part.endswith("*") and len(part) > 1
        content = part[1:-1] if italic else part
        run = paragraph.add_run(content)
        run.italic = italic
        if size:
            run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
        run.font.name = "Calibri"


def add_page_break():
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def add_table(rows):
    header = rows[0]
    body_rows = rows[1:]
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, cell_text in enumerate(header):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(cell_text.strip())
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = INK
    for row in body_rows:
        cells = table.add_row().cells
        for i, cell_text in enumerate(row):
            if i >= len(cells):
                break
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            add_runs_with_italics(p, cell_text.strip(), size=9.5)
    doc.add_paragraph()
    return table


def is_table_row(line):
    return line.strip().startswith("|") and line.strip().endswith("|")


def is_separator_row(line):
    return bool(re.match(r"^\|[\s:\-|]+\|$", line.strip()))


def split_row(line):
    cells = line.strip().strip("|").split("|")
    return [c.strip() for c in cells]


# ---------------------------------------------------------------------------
# Split cover block (before the first standalone "---") from the body
# ---------------------------------------------------------------------------
first_hr = lines.index("---")
cover_lines = [l for l in lines[:first_hr]]
body_lines = lines[first_hr + 1:]

# Cover page
i = 0
while i < len(cover_lines):
    line = cover_lines[i].strip()
    if line.startswith("## "):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line[3:])
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = INK
        p.paragraph_format.space_after = Pt(18)
    elif line.startswith("# "):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line[2:])
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = INK
        p.paragraph_format.space_after = Pt(10)
    elif line:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(12)
        run.font.color.rgb = INK
    i += 1

add_page_break()

# ---------------------------------------------------------------------------
# Body
# ---------------------------------------------------------------------------
i = 0
n = len(body_lines)
in_references = False
while i < n:
    raw = body_lines[i]
    line = raw.strip()

    if not line:
        i += 1
        continue

    if line == "---":
        i += 1
        continue

    if line.startswith("<!--"):
        if "page break" in line.lower():
            add_page_break()
        i += 1
        continue

    if line.startswith("### "):
        text_content = line[4:]
        doc.add_heading(text_content, level=2)
        i += 1
        continue

    if line.startswith("## "):
        text_content = line[3:]
        if text_content.strip().lower() == "references":
            in_references = True
        doc.add_heading(text_content, level=1)
        i += 1
        continue

    if line.startswith("!["):
        m = re.match(r"!\[(.*?)\]\((.*?)\)", line)
        if m:
            alt, relpath = m.group(1), m.group(2)
            img_path = os.path.join(BASE, relpath)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(img_path, width=Inches(6.3))
        i += 1
        continue

    if line.startswith("Source:"):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x52, 0x51, 0x4E)
        p.paragraph_format.space_after = Pt(10)
        i += 1
        continue

    if is_table_row(line):
        table_lines = []
        while i < n and is_table_row(body_lines[i].strip()):
            table_lines.append(body_lines[i].strip())
            i += 1
        rows = [split_row(l) for l in table_lines if not is_separator_row(l)]
        add_table(rows)
        continue

    if line.startswith("- "):
        bullet_text = line[2:]
        p = doc.add_paragraph(style="List Bullet")
        add_runs_with_italics(p, bullet_text)
        i += 1
        continue

    # regular paragraph
    if in_references:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.space_after = Pt(10)
        add_runs_with_italics(p, line)
    else:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        add_runs_with_italics(p, line)
    i += 1

doc.save(OUT)
print("saved:", OUT)
