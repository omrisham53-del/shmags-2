"""
Convert a session-prep markdown plan into a formatted Word .docx.

Usage:
    python plan_to_docx.py "path/to/session_N_plan.md" [output.docx]

If no output path is given, writes alongside the input as <same-name>.docx.
Replaces the old per-session hand-written create_docx.py scripts: it parses the
plan's markdown structure (headings, bold scene labels, bullets, quotes, inline
**bold**/*italic*) so any plan in the standard format converts with one command.

Requires python-docx (pip install python-docx).
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt


def add_inline(paragraph, text):
    """Add text to a paragraph, honoring inline **bold** and *italic* markers."""
    # Split on bold or italic spans while keeping the delimiters.
    pattern = re.compile(r"(\*\*.+?\*\*|\*.+?\*)")
    for chunk in pattern.split(text):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            paragraph.add_run(chunk[2:-2]).bold = True
        elif chunk.startswith("*") and chunk.endswith("*"):
            paragraph.add_run(chunk[1:-1]).italic = True
        else:
            paragraph.add_run(chunk)


def convert(md_path: Path, out_path: Path):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        stripped = line.strip()

        # Horizontal rules / blank lines: skip (Word spacing handles separation).
        if stripped in ("", "---", "***", "___"):
            i += 1
            continue

        # Headings.
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[level:].strip()
            if level == 1:
                doc.add_heading(text, 0)            # Title
            elif level == 2:
                p = doc.add_paragraph()
                run = p.add_run(text)               # Subtitle (e.g. "Adventure 2 - Session 2")
                run.italic = True
                run.font.size = Pt(13)
            else:
                doc.add_heading(text, min(level - 2, 4))  # ### -> H1, #### -> H2, ...
            i += 1
            continue

        # Bullet list items.
        if stripped.startswith(("- ", "* ", "+ ")):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, stripped[2:].strip())
            i += 1
            continue

        # Block quotes (read-aloud / dream text).
        if stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(24)
            run_text = stripped.lstrip(">").strip()
            add_inline(p, run_text)
            for r in p.runs:
                r.italic = True
            i += 1
            continue

        # Everything else: a normal paragraph (handles **Label:** lines too).
        p = doc.add_paragraph()
        add_inline(p, stripped)
        i += 1

    doc.save(str(out_path))
    print(f"Document created: {out_path}")


def main():
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)
    md_path = Path(sys.argv[1])
    if not md_path.exists():
        print(f"Input not found: {md_path}")
        sys.exit(1)
    out_path = Path(sys.argv[2]) if len(sys.argv) > 2 else md_path.with_suffix(".docx")
    convert(md_path, out_path)


if __name__ == "__main__":
    main()
